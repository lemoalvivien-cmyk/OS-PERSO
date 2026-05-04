"""Read both docx files"""
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document

files = [
    r"C:\Users\PC\.openclaw-autoclaw\agents\os-perso\workspace\.autoclaw-attachments\20260504-154557-158828c3-1c1-C--Users-PC-OneDrive-CODE WIFI-B-AUTOCLAW ARTISANAL-CARTOGRAPHIE.docx",
    r"C:\Users\PC\.openclaw-autoclaw\agents\os-perso\workspace\.autoclaw-attachments\20260504-154557-094fd35f-0e0-C--Users-PC-OneDrive-CODE WIFI-B-AUTOCLAW ARTISANAL-CARTOGRAPHIE 2.docx",
]

for fname in files:
    short = fname.split("-")[-1][:40]
    print("=" * 70)
    print(f"FILE: {short}")
    print("=" * 70)
    try:
        doc = Document(fname)
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                print(t)
        for table in doc.tables:
            print("\n[TABLE]")
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                print(" | ".join(cells))
    except Exception as e:
        print(f"ERROR: {e}")
    print()
